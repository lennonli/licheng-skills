#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""企查查资料包批量提取：docx段落目录+全部表格、xlsx/xls全sheet → 单个文本文件。

用法: python3 extract_package.py <资料包目录> [-o 输出文件.md]
"""
import argparse
import sys
from pathlib import Path


def dump_docx(path: Path) -> str:
    import docx
    d = docx.Document(str(path))
    lines = [f"\n######## DOCX: {path.name} ########"]
    lines.append(f"（段落数 {len(d.paragraphs)}，表格数 {len(d.tables)}）")
    lines.append("\n--- 目录/段落（非空，前120条）---")
    n = 0
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t[:100])
            n += 1
            if n >= 120:
                lines.append("……（段落截断）")
                break
    for ti, tbl in enumerate(d.tables):
        lines.append(f"\n--- 表{ti}（{len(tbl.rows)}行 × {len(tbl.columns)}列）---")
        for r in tbl.rows:
            cells = [c.text.strip().replace("\n", "/") for c in r.cells]
            out = []
            for c in cells:  # 相邻去重（合并单元格）
                if not out or out[-1] != c:
                    out.append(c)
            lines.append(" | ".join(out)[:600])
    return "\n".join(lines)


def dump_excel(path: Path) -> str:
    import warnings
    warnings.filterwarnings("ignore")
    import pandas as pd
    lines = [f"\n######## EXCEL: {path.name} ########"]
    try:
        xls = pd.ExcelFile(str(path))
        lines.append(f"Sheets: {xls.sheet_names}")
        for s in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=s, header=None).dropna(how="all")
            lines.append(f"\n--- Sheet [{s}]（{df.shape[0]}行 × {df.shape[1]}列）---")
            for _, row in df.iterrows():
                vals = [str(v).strip() for v in row if str(v) != "nan" and str(v).strip()]
                if vals:
                    lines.append(" | ".join(vals)[:600])
    except Exception as e:
        lines.append(f"读取失败: {e}")
    return "\n".join(lines)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("directory", help="资料包目录")
    ap.add_argument("-o", "--output", default=None, help="输出文件（默认: 目录名_提取.md）")
    args = ap.parse_args()

    root = Path(args.directory).expanduser().resolve()
    if not root.is_dir():
        sys.exit(f"目录不存在: {root}")

    parts = [f"# 资料包提取：{root.name}\n"]
    files = sorted(root.rglob("*"))
    for f in files:
        if not f.is_file() or f.name.startswith("~$"):
            continue
        suf = f.suffix.lower()
        try:
            if suf == ".docx":
                parts.append(dump_docx(f))
            elif suf in (".xlsx", ".xls"):
                parts.append(dump_excel(f))
            elif suf == ".pdf":
                parts.append(f"\n######## PDF（跳过，如需文字层另行pdftotext）: {f.name}")
        except Exception as e:
            parts.append(f"\n!! 解析失败 {f.name}: {e}")

    out = Path(args.output) if args.output else root / f"{root.name}_提取.md"
    out.write_text("\n".join(parts), encoding="utf-8")
    print(f"已输出: {out}")
    print(f"共 {len(parts)-1} 个文件条目，请用 Read 工具分段读取。")


if __name__ == "__main__":
    main()
